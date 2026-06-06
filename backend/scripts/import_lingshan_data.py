# -*- coding: utf-8 -*-
"""从桌面示范资料包导入灵山胜境知识库"""
from pathlib import Path

from docx import Document
import pandas as pd

BASE = Path(r"c:\Users\何煦\Desktop\示范景区公开资料包")
OUT = Path(__file__).resolve().parent.parent / "knowledge_base"


def docx_paragraphs_to_txt(src: Path, dst: Path, title: str):
    doc = Document(src)
    lines = [f"# {title}\n"]
    lines.extend(p.text for p in doc.paragraphs if p.text.strip())
    for ti, table in enumerate(doc.tables):
        lines.append(f"\n## 数据表 {ti + 1}\n")
        for row in table.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))
    dst.write_text("\n".join(lines), encoding="utf-8")
    print(f"  wrote {dst.name} ({dst.stat().st_size} bytes)")


def import_xlsx_behavior(src: Path, dst: Path):
    df = pd.read_excel(src)
    keywords = ["灵山", "灵山大佛", "梵宫", "九龙", "拈花", "胜境", "五印", "祥符"]
    mask = df["attraction_name"].astype(str).apply(
        lambda x: any(k in x for k in keywords)
    )
    ls = df[mask]
    lines = [
        "# 灵山胜境 游客行为与体验数据摘要",
        "",
        "以下数据来源于示范景区旅游行为分析数据集，供导览推荐与运营分析参考。",
        f"灵山相关记录数：{len(ls)} 条",
        "",
    ]
    if ls.empty:
        lines.append("（数据集中未匹配到灵山关键词，已保留全库热门景区统计供参考）")
        top = df.groupby("attraction_name").size().sort_values(ascending=False).head(20)
        lines.append("\n## 数据集热门景区 TOP20（参考）\n")
        for name, cnt in top.items():
            lines.append(f"- {name}：{cnt} 条访问记录")
    else:
        agg = ls.groupby("attraction_name").agg(
            attraction_content=("attraction_content", "first"),
            attraction_type=("attraction_type", "first"),
            satisfaction=("satisfaction", "mean"),
            stay_duration=("stay_duration", "mean"),
            total_cost=("total_cost", "mean"),
        ).reset_index()
        for _, r in agg.iterrows():
            lines.append(f"\n## {r['attraction_name']}\n")
            lines.append(f"景点类型：{r['attraction_type']}")
            lines.append(f"平均停留：{r['stay_duration']:.0f} 分钟")
            lines.append(f"平均消费：{r['total_cost']:.0f} 元")
            lines.append(f"平均满意度：{r['satisfaction']:.1f} / 5")
            content = str(r["attraction_content"])[:4000]
            lines.append(f"\n{content}\n")

    # 全库统计：热门问题、消费结构（灵山景区演示用）
    lines.append("\n## 长三角景区游客行为洞察（全库统计）\n")
    lines.append(f"- 总样本量：{len(df)} 条")
    lines.append(f"- 覆盖景点数：{df['attraction_name'].nunique()} 个")
    avg_sat = df["satisfaction"].mean()
    lines.append(f"- 全库平均满意度：{avg_sat:.2f} / 5")
    top_names = df["attraction_name"].value_counts().head(15)
    lines.append("\n### 访问量较高的景区（TOP15）\n")
    for name, cnt in top_names.items():
        lines.append(f"- {name}：{cnt} 次")

    dst.write_text("\n".join(lines), encoding="utf-8")
    print(f"  wrote {dst.name} ({dst.stat().st_size} bytes), lingshan rows={len(ls)}")


def main():
    OUT.mkdir(parents=True, exist_ok=True)

    print("Importing from:", BASE)
    docx_paragraphs_to_txt(
        BASE / "灵山胜境：历史、文化、景点特色与个性化游览指南.docx",
        OUT / "灵山胜境_游览指南.txt",
        "灵山胜境：历史、文化、景点特色与个性化游览指南",
    )
    docx_paragraphs_to_txt(
        BASE / "灵山胜境 景点结构化数据集.docx",
        OUT / "灵山胜境_景点结构化.txt",
        "灵山胜境 景点结构化数据集",
    )
    import_xlsx_behavior(
        BASE / "景点景区旅游数据行为分析数据.xlsx",
        OUT / "灵山胜境_游客行为数据.txt",
    )

    # 移除旧示例文件（可选）
    for old in ["景区介绍.txt", "历史文化.txt"]:
        p = OUT / old
        if p.exists():
            p.unlink()
            print(f"  removed old {old}")

    print("Done. Run: python rag_chain.py to rebuild chroma_db")


if __name__ == "__main__":
    main()
